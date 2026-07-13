from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path

import fitz
from docx import Document

from resume_matcher.application.dto import ParsedDocument
from resume_matcher.domain.exceptions import (
    DocumentExtractionError,
    DocumentTooLargeError,
    InvalidDocumentError,
    UnsupportedDocumentError,
)

PDF_MIME_TYPES = {"application/pdf", "application/x-pdf"}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
    "application/zip",
}


class SecureDocumentParser:
    def __init__(
        self,
        *,
        max_bytes: int,
        max_pages: int = 100,
        max_docx_uncompressed_bytes: int = 50 * 1024 * 1024,
        max_text_chars: int = 500_000,
    ) -> None:
        self._max_bytes = max_bytes
        self._max_pages = max_pages
        self._max_docx_uncompressed_bytes = max_docx_uncompressed_bytes
        self._max_text_chars = max_text_chars

    def parse(self, *, filename: str, content_type: str | None, content: bytes) -> ParsedDocument:
        if not content:
            raise InvalidDocumentError("The uploaded file is empty")
        if len(content) > self._max_bytes:
            raise DocumentTooLargeError(
                f"The uploaded file exceeds the {self._max_bytes}-byte limit"
            )
        extension = Path(filename).suffix.casefold()
        if extension == ".pdf":
            self._validate_content_type(content_type, PDF_MIME_TYPES)
            if not content.startswith(b"%PDF-"):
                raise InvalidDocumentError("The file extension is PDF but its signature is invalid")
            return self._parse_pdf(content)
        if extension == ".docx":
            self._validate_content_type(content_type, DOCX_MIME_TYPES)
            if not content.startswith(b"PK"):
                raise InvalidDocumentError(
                    "The file extension is DOCX but its signature is invalid"
                )
            self._validate_docx_archive(content)
            return self._parse_docx(content)
        raise UnsupportedDocumentError("Only .pdf and .docx resume files are supported")

    @staticmethod
    def _validate_content_type(content_type: str | None, allowed: set[str]) -> None:
        if content_type and content_type.casefold() not in allowed:
            raise InvalidDocumentError(
                f"Declared content type '{content_type}' does not match the file extension"
            )

    def _parse_pdf(self, content: bytes) -> ParsedDocument:
        try:
            with fitz.open(stream=content, filetype="pdf") as document:
                if document.needs_pass:
                    raise InvalidDocumentError("Password-protected PDFs are not supported")
                if document.page_count > self._max_pages:
                    raise InvalidDocumentError(
                        f"PDF has {document.page_count} pages; maximum is {self._max_pages}"
                    )
                text = "\n".join(page.get_text("text", sort=True) for page in document)
        except InvalidDocumentError:
            raise
        except Exception as exc:
            raise DocumentExtractionError("Unable to safely extract text from the PDF") from exc
        return self._finish(text, source="PDF")

    def _validate_docx_archive(self, content: bytes) -> None:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
                required = {"[Content_Types].xml", "word/document.xml"}
                if not required.issubset(names):
                    raise InvalidDocumentError("The DOCX package is missing required Word parts")
                if any(name.casefold().endswith("vbaproject.bin") for name in names):
                    raise InvalidDocumentError("Macro-enabled Word packages are not supported")
                total_uncompressed = sum(item.file_size for item in archive.infolist())
                if total_uncompressed > self._max_docx_uncompressed_bytes:
                    raise InvalidDocumentError("The DOCX archive expands beyond the safety limit")
                for item in archive.infolist():
                    if item.file_size > 1_000_000 and item.compress_size > 0:
                        ratio = item.file_size / item.compress_size
                        if ratio > 100:
                            raise InvalidDocumentError(
                                "The DOCX archive has a suspicious compression ratio"
                            )
        except InvalidDocumentError:
            raise
        except (zipfile.BadZipFile, OSError) as exc:
            raise InvalidDocumentError("The DOCX package is corrupted") from exc

    def _parse_docx(self, content: bytes) -> ParsedDocument:
        try:
            document = Document(io.BytesIO(content))
            blocks: list[str] = []
            blocks.extend(
                paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
            )
            for table in document.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        blocks.append(" | ".join(values))
            text = "\n".join(blocks)
        except Exception as exc:
            raise DocumentExtractionError(
                "Unable to safely extract text from the DOCX file"
            ) from exc
        return self._finish(text, source="DOCX")

    def _finish(self, text: str, *, source: str) -> ParsedDocument:
        normalized = text.replace("\x00", "")
        normalized = re.sub(r"[ \t]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized).strip()
        if not normalized:
            raise DocumentExtractionError(
                f"No selectable text was found in the {source}; scanned documents require OCR"
            )
        warnings: list[str] = []
        if len(normalized) > self._max_text_chars:
            normalized = normalized[: self._max_text_chars]
            warnings.append(
                f"Extracted text was truncated to {self._max_text_chars} characters for safety."
            )
        return ParsedDocument(text=normalized, warnings=tuple(warnings))
