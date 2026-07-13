from __future__ import annotations

import io
import json
from dataclasses import asdict
from datetime import UTC, datetime
from typing import Any, ClassVar
from xml.sax.saxutils import escape

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Spacer,
)

from resume_matcher.application.dto import ExportArtifact
from resume_matcher.domain.entities import (
    JobProfile,
    MatchAnalysis,
    OptimizedResume,
    ResumeProfile,
)

ResumeLike = ResumeProfile | OptimizedResume
_NAVY = RGBColor(31, 78, 121)


class MultiFormatResumeExporter:
    SUPPORTED_FORMATS: ClassVar[frozenset[str]] = frozenset({"json", "docx", "pdf"})

    def export(
        self,
        *,
        format_name: str,
        resume: ResumeLike,
        match: MatchAnalysis,
        job: JobProfile,
    ) -> ExportArtifact:
        normalized = format_name.casefold()
        if normalized not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported export format '{format_name}'; use json, docx, or pdf")
        stem = "optimized-resume" if match.optimized_resume else "resume-analysis"
        if normalized == "json":
            content = self._json(resume, match, job)
            media_type = "application/json"
        elif normalized == "docx":
            content = self._docx(resume)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            content = self._pdf(resume)
            media_type = "application/pdf"
        return ExportArtifact(
            content=content,
            media_type=media_type,
            filename=f"{stem}-{match.id}.{normalized}",
        )

    @staticmethod
    def _json(resume: ResumeLike, match: MatchAnalysis, job: JobProfile) -> bytes:
        payload: dict[str, Any] = {
            "schema_version": "1.0.0",
            "generated_at": datetime.now(UTC).isoformat(),
            "resume": asdict(resume),
            "target_job": asdict(job),
            "match_analysis": asdict(match.result),
            "optimized": match.optimized_resume is not None,
        }
        return json.dumps(payload, indent=2, ensure_ascii=False, default=str).encode("utf-8")

    def _docx(self, resume: ResumeLike) -> bytes:
        document = Document()
        section = document.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)

        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(9.5)
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.05

        if "Resume Section" not in styles:
            section_style = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
        else:
            section_style = styles["Resume Section"]
        section_style.font.name = "Arial"
        section_style.font.size = Pt(11)
        section_style.font.bold = True
        section_style.font.color.rgb = _NAVY
        section_style.paragraph_format.space_before = Pt(8)
        section_style.paragraph_format.space_after = Pt(3)

        name = document.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.paragraph_format.space_after = Pt(2)
        name_run = name.add_run(resume.name or "Candidate")
        name_run.font.name = "Arial"
        name_run.font.size = Pt(21)
        name_run.font.bold = True
        name_run.font.color.rgb = _NAVY

        if resume.headline:
            headline = document.add_paragraph()
            headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
            headline.paragraph_format.space_after = Pt(3)
            run = headline.add_run(resume.headline)
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
            run.font.bold = True

        contact_values = [value for value in (resume.location, resume.email, resume.phone) if value]
        if contact_values:
            contact = document.add_paragraph()
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.paragraph_format.space_after = Pt(7)
            run = contact.add_run(" | ".join(contact_values))
            run.font.name = "Arial"
            run.font.size = Pt(8.5)

        if resume.summary:
            self._docx_heading(document, "PROFESSIONAL SUMMARY")
            document.add_paragraph(resume.summary)

        if resume.skills:
            self._docx_heading(document, "CORE SKILLS")
            document.add_paragraph(" | ".join(skill.name for skill in resume.skills))

        if resume.experiences:
            self._docx_heading(document, "EXPERIENCE")
            for experience in resume.experiences:
                role = document.add_paragraph()
                role.paragraph_format.keep_with_next = True
                role.paragraph_format.space_before = Pt(4)
                title_run = role.add_run(experience.title)
                title_run.bold = True
                title_run.font.size = Pt(10)
                company_run = role.add_run(f" | {experience.company}")
                company_run.italic = True
                dates = " - ".join(
                    value for value in (experience.start_date, experience.end_date) if value
                )
                details = " | ".join(value for value in (dates, experience.location) if value)
                if details:
                    detail = document.add_paragraph(details)
                    detail.paragraph_format.keep_with_next = bool(experience.bullets)
                    detail.paragraph_format.space_after = Pt(2)
                    detail.runs[0].font.size = Pt(8.5)
                    detail.runs[0].font.color.rgb = RGBColor(89, 89, 89)
                for bullet in experience.bullets:
                    paragraph = document.add_paragraph(style="List Bullet")
                    paragraph.paragraph_format.left_indent = Inches(0.18)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
                    paragraph.paragraph_format.space_after = Pt(1.5)
                    paragraph.add_run(bullet)

        if resume.education:
            self._docx_heading(document, "EDUCATION")
            for education in resume.education:
                paragraph = document.add_paragraph()
                degree = paragraph.add_run(education.degree)
                degree.bold = True
                suffix = f" | {education.institution}"
                if education.field:
                    suffix = f" in {education.field}{suffix}"
                if education.graduation_year:
                    suffix += f" | {education.graduation_year}"
                paragraph.add_run(suffix)

        if resume.certifications:
            self._docx_heading(document, "CERTIFICATIONS")
            for certification in resume.certifications:
                document.add_paragraph(certification, style="List Bullet")

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _docx_heading(document: DocxDocument, title: str) -> None:
        paragraph = document.add_paragraph(title, style="Resume Section")
        paragraph.paragraph_format.keep_with_next = True
        properties = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "1F4E79")
        borders.append(bottom)
        properties.append(borders)

    def _pdf(self, resume: ResumeLike) -> bytes:
        buffer = io.BytesIO()
        document = BaseDocTemplate(
            buffer,
            pagesize=LETTER,
            leftMargin=0.72 * inch,
            rightMargin=0.72 * inch,
            topMargin=0.55 * inch,
            bottomMargin=0.55 * inch,
            title=f"Resume - {resume.name or 'Candidate'}",
            author=resume.name or "Candidate",
        )
        frame = Frame(
            document.leftMargin,
            document.bottomMargin,
            document.width,
            document.height,
            id="resume",
        )
        document.addPageTemplates(
            [PageTemplate(id="resume", frames=[frame], onPage=self._pdf_footer)]
        )
        styles = getSampleStyleSheet()
        name_style = ParagraphStyle(
            "ResumeName",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=22,
            textColor=colors.HexColor("#1F4E79"),
            alignment=TA_CENTER,
            spaceAfter=3,
        )
        headline_style = ParagraphStyle(
            "Headline",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=2,
        )
        contact_style = ParagraphStyle(
            "Contact",
            parent=styles["Normal"],
            fontSize=8.5,
            leading=10,
            alignment=TA_CENTER,
            spaceAfter=7,
        )
        heading_style = ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=12,
            textColor=colors.HexColor("#1F4E79"),
            borderColor=colors.HexColor("#1F4E79"),
            borderWidth=0,
            borderPadding=0,
            spaceBefore=7,
            spaceAfter=3,
        )
        body_style = ParagraphStyle(
            "ResumeBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=11.2,
            spaceAfter=3,
            alignment=TA_LEFT,
        )
        bullet_style = ParagraphStyle(
            "ResumeBullet",
            parent=body_style,
            leftIndent=12,
            firstLineIndent=-8,
            bulletIndent=2,
            spaceAfter=1.5,
        )
        story: list[Any] = [Paragraph(escape(resume.name or "Candidate"), name_style)]
        if resume.headline:
            story.append(Paragraph(escape(resume.headline), headline_style))
        contact = [value for value in (resume.location, resume.email, resume.phone) if value]
        if contact:
            story.append(Paragraph(escape(" | ".join(contact)), contact_style))
        if resume.summary:
            story.extend(
                [
                    Paragraph("PROFESSIONAL SUMMARY", heading_style),
                    Paragraph(escape(resume.summary), body_style),
                ]
            )
        if resume.skills:
            story.extend(
                [
                    Paragraph("CORE SKILLS", heading_style),
                    Paragraph(
                        escape(" | ".join(skill.name for skill in resume.skills)), body_style
                    ),
                ]
            )
        if resume.experiences:
            story.append(Paragraph("EXPERIENCE", heading_style))
            for experience in resume.experiences:
                role = f"<b>{escape(experience.title)}</b> | <i>{escape(experience.company)}</i>"
                story.append(Paragraph(role, body_style))
                details = " | ".join(
                    value
                    for value in (
                        " - ".join(
                            value for value in (experience.start_date, experience.end_date) if value
                        ),
                        experience.location or "",
                    )
                    if value
                )
                if details:
                    story.append(Paragraph(escape(details), body_style))
                for bullet in experience.bullets:
                    story.append(Paragraph(escape(bullet), bullet_style, bulletText="-"))
                story.append(Spacer(1, 2))
        if resume.education:
            story.append(Paragraph("EDUCATION", heading_style))
            for education in resume.education:
                text = f"<b>{escape(education.degree)}</b>"
                if education.field:
                    text += f" in {escape(education.field)}"
                text += f" | {escape(education.institution)}"
                if education.graduation_year:
                    text += f" | {education.graduation_year}"
                story.append(Paragraph(text, body_style))
        if resume.certifications:
            story.append(Paragraph("CERTIFICATIONS", heading_style))
            for certification in resume.certifications:
                story.append(Paragraph(escape(certification), bullet_style, bulletText="-"))
        document.build(story)
        return buffer.getvalue()

    @staticmethod
    def _pdf_footer(canvas: Any, document: BaseDocTemplate) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawCentredString(LETTER[0] / 2, 0.32 * inch, f"Page {document.page}")
        canvas.restoreState()
