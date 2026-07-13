from __future__ import annotations

import io
import zipfile

import fitz
import pytest
from docx import Document

from resume_matcher.domain.exceptions import (
    DocumentExtractionError,
    DocumentTooLargeError,
    InvalidDocumentError,
    UnsupportedDocumentError,
)
from resume_matcher.infrastructure.parsing.document_parser import SecureDocumentParser

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _pdf_bytes(text: str | None = "Jane Doe\nPython engineer", *, pages: int = 1) -> bytes:
    document = fitz.open()
    try:
        for _ in range(pages):
            page = document.new_page()
            if text is not None:
                page.insert_text((72, 72), text)
        return document.tobytes()
    finally:
        document.close()


def _encrypted_pdf_bytes() -> bytes:
    document = fitz.open()
    try:
        document.new_page().insert_text((72, 72), "Private resume")
        return document.tobytes(
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw="owner-password",
            user_pw="user-password",
        )
    finally:
        document.close()


def _docx_bytes(*, text: str | None = "Jane Doe", table: bool = False) -> bytes:
    document = Document()
    if text is not None:
        document.add_paragraph(text)
    if table:
        row = document.add_table(rows=1, cols=2).rows[0]
        row.cells[0].text = "Skill"
        row.cells[1].text = "Python"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _archive_bytes(
    *,
    include_required: bool = True,
    extra: dict[str, bytes] | None = None,
    compression: int = zipfile.ZIP_DEFLATED,
) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=compression) as archive:
        if include_required:
            archive.writestr("[Content_Types].xml", b"<Types />")
            archive.writestr("word/document.xml", b"<document />")
        for name, value in (extra or {}).items():
            archive.writestr(name, value)
    return buffer.getvalue()


def test_parse_pdf_extracts_and_normalizes_text() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)

    result = parser.parse(
        filename="RESUME.PDF",
        content_type=PDF_CONTENT_TYPE,
        content=_pdf_bytes("Jane   Doe\n\nPython engineer"),
    )

    assert result.text == "Jane Doe\n\nPython engineer"
    assert result.warnings == ()


def test_parse_docx_extracts_paragraphs_and_tables() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)

    result = parser.parse(
        filename="resume.docx",
        content_type="application/octet-stream",
        content=_docx_bytes(text="Jane Doe\nPlatform Engineer", table=True),
    )

    assert result.text == "Jane Doe\nPlatform Engineer\nSkill | Python"
    assert result.warnings == ()


def test_parse_accepts_an_absent_content_type() -> None:
    result = SecureDocumentParser(max_bytes=1_000_000).parse(
        filename="resume.pdf", content_type=None, content=_pdf_bytes()
    )

    assert "Jane Doe" in result.text


@pytest.mark.parametrize(
    ("filename", "content_type", "content", "exception", "message"),
    [
        ("resume.pdf", PDF_CONTENT_TYPE, b"", InvalidDocumentError, "empty"),
        (
            "resume.txt",
            "text/plain",
            b"plain text",
            UnsupportedDocumentError,
            "Only .pdf and .docx",
        ),
        (
            "resume.pdf",
            "text/plain",
            b"%PDF-content",
            InvalidDocumentError,
            "does not match",
        ),
        (
            "resume.docx",
            "text/plain",
            b"PK-content",
            InvalidDocumentError,
            "does not match",
        ),
        (
            "resume.pdf",
            PDF_CONTENT_TYPE,
            b"not a pdf",
            InvalidDocumentError,
            "signature is invalid",
        ),
        (
            "resume.docx",
            DOCX_CONTENT_TYPE,
            b"not a docx",
            InvalidDocumentError,
            "signature is invalid",
        ),
    ],
)
def test_parse_rejects_invalid_upload_boundaries(
    filename: str,
    content_type: str,
    content: bytes,
    exception: type[Exception],
    message: str,
) -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)

    with pytest.raises(exception, match=message):
        parser.parse(filename=filename, content_type=content_type, content=content)


def test_parse_rejects_file_over_byte_limit_before_processing() -> None:
    parser = SecureDocumentParser(max_bytes=10)

    with pytest.raises(DocumentTooLargeError, match="10-byte limit"):
        parser.parse(filename="resume.pdf", content_type=PDF_CONTENT_TYPE, content=b"%PDF-" * 3)


def test_parse_wraps_malformed_pdf_errors() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)

    with pytest.raises(DocumentExtractionError, match="Unable to safely extract"):
        parser.parse(
            filename="resume.pdf", content_type=PDF_CONTENT_TYPE, content=b"%PDF-not-really-pdf"
        )


def test_parse_rejects_password_protected_pdf() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)

    with pytest.raises(InvalidDocumentError, match="Password-protected"):
        parser.parse(
            filename="resume.pdf",
            content_type=PDF_CONTENT_TYPE,
            content=_encrypted_pdf_bytes(),
        )


def test_parse_rejects_pdf_over_page_limit() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000, max_pages=1)

    with pytest.raises(InvalidDocumentError, match="2 pages; maximum is 1"):
        parser.parse(
            filename="resume.pdf", content_type=PDF_CONTENT_TYPE, content=_pdf_bytes(pages=2)
        )


def test_parse_rejects_pdf_without_selectable_text() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)

    with pytest.raises(DocumentExtractionError, match="scanned documents require OCR"):
        parser.parse(
            filename="resume.pdf", content_type=PDF_CONTENT_TYPE, content=_pdf_bytes(text=None)
        )


def test_parse_rejects_corrupt_docx_archive() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)

    with pytest.raises(InvalidDocumentError, match="package is corrupted"):
        parser.parse(filename="resume.docx", content_type=DOCX_CONTENT_TYPE, content=b"PKbroken")


def test_parse_rejects_docx_missing_required_word_parts() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)
    content = _archive_bytes(include_required=False, extra={"unrelated.txt": b"data"})

    with pytest.raises(InvalidDocumentError, match="missing required Word parts"):
        parser.parse(filename="resume.docx", content_type=DOCX_CONTENT_TYPE, content=content)


def test_parse_rejects_macro_enabled_docx() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)
    content = _archive_bytes(extra={"word/vbaProject.bin": b"macro"})

    with pytest.raises(InvalidDocumentError, match="Macro-enabled"):
        parser.parse(filename="resume.docx", content_type=DOCX_CONTENT_TYPE, content=content)


def test_parse_rejects_docx_over_uncompressed_size_limit() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000, max_docx_uncompressed_bytes=100)
    content = _archive_bytes(extra={"large.txt": b"x" * 101})

    with pytest.raises(InvalidDocumentError, match="expands beyond the safety limit"):
        parser.parse(filename="resume.docx", content_type=DOCX_CONTENT_TYPE, content=content)


def test_parse_rejects_docx_with_suspicious_compression_ratio() -> None:
    parser = SecureDocumentParser(
        max_bytes=1_000_000,
        max_docx_uncompressed_bytes=2_000_000,
    )
    content = _archive_bytes(extra={"word/huge.xml": b"0" * 1_000_001})

    with pytest.raises(InvalidDocumentError, match="suspicious compression ratio"):
        parser.parse(filename="resume.docx", content_type=DOCX_CONTENT_TYPE, content=content)


def test_parse_rejects_docx_without_extractable_text() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000)

    with pytest.raises(DocumentExtractionError, match="scanned documents require OCR"):
        parser.parse(
            filename="resume.docx",
            content_type=DOCX_CONTENT_TYPE,
            content=_docx_bytes(text=None),
        )


def test_parse_truncates_excess_extracted_text_with_warning() -> None:
    parser = SecureDocumentParser(max_bytes=1_000_000, max_text_chars=12)

    result = parser.parse(
        filename="resume.docx",
        content_type=DOCX_CONTENT_TYPE,
        content=_docx_bytes(text="0123456789ABCDEFGHIJ"),
    )

    assert result.text == "0123456789AB"
    assert result.warnings == ("Extracted text was truncated to 12 characters for safety.",)
